"""File analysis tools for various file types."""

import base64
from pathlib import Path
from typing import Any

from core.tools.base import Tool, ToolExecutionError


class ReadImageTool(Tool):
    """Tool for reading and analyzing image files with optional OCR."""

    @property
    def name(self) -> str:
        return "read_image"

    @property
    def description(self) -> str:
        return "读取图片文件并分析内容，支持获取图片信息和OCR文字识别"

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "图片文件路径"
                },
                "enable_ocr": {
                    "type": "boolean",
                    "description": "是否启用OCR文字识别",
                    "default": False
                }
            },
            "required": ["file_path"]
        }

    async def execute(self, **kwargs: Any) -> str:
        """Execute the image reading tool."""
        file_path = kwargs.get("file_path")
        enable_ocr = kwargs.get("enable_ocr", False)

        if not file_path:
            raise ToolExecutionError("file_path is required")

        path = Path(file_path)
        if not path.exists():
            raise ToolExecutionError(f"File not found: {file_path}")

        try:
            # Try to import PIL
            try:
                from PIL import Image
            except ImportError:
                raise ToolExecutionError(
                    "PIL/Pillow is not installed. Install it with: pip install pillow"
                )

            with Image.open(path) as img:
                result = [
                    f"图片信息:",
                    f"  文件名: {path.name}",
                    f"  格式: {img.format}",
                    f"  模式: {img.mode}",
                    f"  尺寸: {img.width} x {img.height} 像素",
                    f"  大小: {path.stat().st_size} 字节",
                ]

                # Get EXIF data if available
                try:
                    exif = img._getexif()
                    if exif:
                        result.append("  包含EXIF元数据")
                except (AttributeError, TypeError):
                    pass

                # Perform OCR if requested
                if enable_ocr:
                    try:
                        import pytesseract
                        text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                        if text.strip():
                            result.append(f"\nOCR识别文字:\n{text}")
                        else:
                            result.append("\nOCR未检测到文字内容")
                    except ImportError:
                        result.append("\n注意: pytesseract未安装，无法进行OCR识别")
                    except Exception as e:
                        result.append(f"\nOCR识别失败: {e}")

                # Generate base64 for preview (small images only)
                if img.width * img.height <= 500000:  # < 500k pixels
                    import io
                    buffer = io.BytesIO()
                    img.save(buffer, format='PNG')
                    img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                    result.append(f"\nBase64预览: {img_base64[:100]}...")

                return "\n".join(result)

        except Exception as e:
            raise ToolExecutionError(f"Failed to read image: {e}")


class ReadDocumentTool(Tool):
    """Tool for reading document files (PDF, DOCX, TXT)."""

    @property
    def name(self) -> str:
        return "read_document"

    @property
    def description(self) -> str:
        return "读取文档文件（PDF、DOCX、TXT）并提取内容"

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "文档文件路径"
                },
                "max_pages": {
                    "type": "integer",
                    "description": "最大读取页数（PDF）",
                    "default": 50
                }
            },
            "required": ["file_path"]
        }

    async def execute(self, **kwargs: Any) -> str:
        """Execute the document reading tool."""
        file_path = kwargs.get("file_path")
        max_pages = kwargs.get("max_pages", 50)

        if not file_path:
            raise ToolExecutionError("file_path is required")

        path = Path(file_path)
        if not path.exists():
            raise ToolExecutionError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        try:
            if ext == ".pdf":
                return await self._read_pdf(path, max_pages)
            elif ext in [".docx", ".doc"]:
                return await self._read_docx(path)
            elif ext == ".txt":
                return await self._read_txt(path)
            else:
                raise ToolExecutionError(f"Unsupported document format: {ext}")

        except Exception as e:
            raise ToolExecutionError(f"Failed to read document: {e}")

    async def _read_pdf(self, path: Path, max_pages: int) -> str:
        """Read PDF file using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            raise ToolExecutionError(
                "pdfplumber is not installed. Install it with: pip install pdfplumber"
            )

        result = [f"PDF文档: {path.name}"]
        pages_text = []
        tables_data = []

        with pdfplumber.open(path) as pdf:
            result.append(f"总页数: {len(pdf.pages)}")

            for i, page in enumerate(pdf.pages[:max_pages]):
                # Extract text
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"--- 第 {i+1} 页 ---\n{text}")

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    for j, table in enumerate(tables):
                        if table and any(row for row in table):
                            tables_data.append(f"\n--- 表格 {len(tables_data)+1} (第{i+1}页) ---")
                            for row in table:
                                row_str = " | ".join(str(cell or "") for cell in row)
                                tables_data.append(row_str)

        if pages_text:
            result.append("\n文本内容:\n" + "\n\n".join(pages_text[:max_pages]))

        if tables_data:
            result.append("\n表格数据:\n" + "\n".join(tables_data))

        return "\n".join(result)

    async def _read_docx(self, path: Path) -> str:
        """Read DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ToolExecutionError(
                "python-docx is not installed. Install it with: pip install python-docx"
            )

        result = [f"Word文档: {path.name}"]
        doc = Document(path)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        if paragraphs:
            result.append("\n段落内容:\n" + "\n\n".join(paragraphs))

        # Extract tables
        tables_data = []
        for i, table in enumerate(doc.tables):
            tables_data.append(f"\n--- 表格 {i+1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_data.append(" | ".join(cells))

        if tables_data:
            result.append("\n表格数据:\n" + "\n".join(tables_data))

        return "\n".join(result)

    async def _read_txt(self, path: Path) -> str:
        """Read plain text file."""
        encoding = 'utf-8'
        content = ""

        # Try different encodings
        for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                content = path.read_text(encoding=enc)
                encoding = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue

        result = [
            f"文本文件: {path.name}",
            f"编码: {encoding}",
            f"大小: {len(content)} 字符",
            f"\n内容:\n{content}"
        ]

        return "\n".join(result)


class ParseTableTool(Tool):
    """Tool for parsing spreadsheet files (CSV, XLSX)."""

    @property
    def name(self) -> str:
        return "parse_table"

    @property
    def description(self) -> str:
        return "解析表格文件（CSV、XLSX）并返回结构化数据"

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "表格文件路径"
                },
                "head_rows": {
                    "type": "integer",
                    "description": "显示的行数",
                    "default": 10
                },
                "encoding": {
                    "type": "string",
                    "description": "文件编码（仅CSV）",
                    "default": "utf-8"
                }
            },
            "required": ["file_path"]
        }

    async def execute(self, **kwargs: Any) -> str:
        """Execute the table parsing tool."""
        file_path = kwargs.get("file_path")
        head_rows = kwargs.get("head_rows", 10)
        encoding = kwargs.get("encoding", "utf-8")

        if not file_path:
            raise ToolExecutionError("file_path is required")

        path = Path(file_path)
        if not path.exists():
            raise ToolExecutionError(f"File not found: {file_path}")

        try:
            # Try to import pandas
            try:
                import pandas as pd
            except ImportError:
                raise ToolExecutionError(
                    "pandas is not installed. Install it with: pip install pandas"
                )

            ext = path.suffix.lower()

            # Read the file
            if ext == ".csv":
                df = pd.read_csv(path, encoding=encoding)
            elif ext in [".xlsx", ".xls"]:
                df = pd.read_excel(path, engine='openpyxl')
            else:
                raise ToolExecutionError(f"Unsupported table format: {ext}")

            # Build result
            result = [
                f"表格文件: {path.name}",
                f"格式: {ext[1:].upper()}",
                f"行数: {len(df)}",
                f"列数: {len(df.columns)}",
                f"\n列名: {', '.join(df.columns.tolist())}",
                f"\n数据类型:",
            ]

            for col, dtype in df.dtypes.items():
                result.append(f"  {col}: {dtype}")

            result.append(f"\n统计信息:")
            result.append(df.describe().to_string())

            result.append(f"\n前{min(head_rows, len(df))}行数据:")
            result.append(df.head(head_rows).to_string(index=False))

            return "\n".join(result)

        except Exception as e:
            raise ToolExecutionError(f"Failed to parse table: {e}")